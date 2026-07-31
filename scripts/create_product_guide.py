from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Windows\Documents\INDIVIDUAL PROJECT\hanziwork")
QA = ROOT / "docs" / "design-qa"
ASSETS = ROOT / "docs" / "word-assets"
OUT = ROOT / "docs" / "HanziWork_Blueprint_v0.2.docx"
ASSETS.mkdir(parents=True, exist_ok=True)

FOREST = "0F6A58"
FOREST_DARK = "173A34"
MINT = "DCEBE2"
MINT_LIGHT = "EFF6F1"
PAPER = "F7F6F1"
GOLD = "C5903F"
GOLD_LIGHT = "FFF5DD"
LINE = "DCE5DF"
WHITE = "FFFFFF"
INK_SOFT = "536A64"
BLUE = "315B7D"

FONT_REG = r"C:\Windows\Fonts\segoeui.ttf"
FONT_SEMI = r"C:\Windows\Fonts\seguisb.ttf"
FONT_BOLD = r"C:\Windows\Fonts\segoeuib.ttf"


def pil_font(size: int, bold: bool = False, semi: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_BOLD if bold else FONT_SEMI if semi else FONT_REG
    return ImageFont.truetype(path, size)


def fit_image(image: Image.Image, box: tuple[int, int], contain: bool = True) -> Image.Image:
    image = image.convert("RGB")
    if contain:
        image.thumbnail(box, Image.Resampling.LANCZOS)
        return image
    return ImageOps.fit(image, box, method=Image.Resampling.LANCZOS)


def draw_wrapped(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font: ImageFont.FreeTypeFont,
                 fill: str, max_width: int, line_gap: int = 6, max_lines: int | None = None) -> int:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    if max_lines and len(lines) > max_lines:
        lines = lines[:max_lines]
        tail = lines[-1]
        while draw.textbbox((0, 0), tail + "…", font=font)[2] > max_width and tail:
            tail = tail[:-1]
        lines[-1] = tail + "…"
    x, y = xy
    line_height = font.size + line_gap
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height
    return y


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = FOREST,
          width: int = 5) -> None:
    draw.line([start, end], fill=f"#{color}", width=width)
    x2, y2 = end
    x1, y1 = start
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - sign * 18, y2 - 10), (x2 - sign * 18, y2 + 10)]
    else:
        sign = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 10, y2 - sign * 18), (x2 + 10, y2 - sign * 18)]
    draw.polygon(points, fill=f"#{color}")


def rounded_task(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, subtitle: str = "",
                 fill: str = WHITE, outline: str = LINE, accent: str | None = None) -> None:
    draw.rounded_rectangle(box, radius=22, fill=f"#{fill}", outline=f"#{outline}", width=3)
    x1, y1, x2, y2 = box
    if accent:
        draw.rounded_rectangle((x1, y1, x1 + 12, y2), radius=8, fill=f"#{accent}")
    tx = x1 + 24
    draw_wrapped(draw, (tx, y1 + 18), title, pil_font(25, semi=True), f"#{FOREST_DARK}", x2 - tx - 18, max_lines=2)
    if subtitle:
        draw_wrapped(draw, (tx, y1 + 62), subtitle, pil_font(18), f"#{INK_SOFT}", x2 - tx - 18, max_lines=2)


def build_responsive_overview() -> Path:
    laptop = Image.open(QA / "home-laptop.png").convert("RGB")
    phone = Image.open(QA / "home-mobile-crop.png").convert("RGB")
    canvas = Image.new("RGB", (1800, 1080), f"#{PAPER}")
    draw = ImageDraw.Draw(canvas)
    draw.text((80, 52), "PHƯƠNG ÁN 3 · TRẠM HỌC TẬP RESPONSIVE", font=pil_font(25, bold=True), fill=f"#{FOREST}")
    draw.text((80, 92), "Một ưu tiên rõ ràng trên laptop, một cột dễ chạm trên điện thoại.", font=pil_font(38, semi=True), fill=f"#{FOREST_DARK}")
    laptop_fit = fit_image(laptop, (1240, 720))
    lx, ly = 80, 200
    draw.rounded_rectangle((lx - 18, ly - 18, lx + laptop_fit.width + 18, ly + laptop_fit.height + 18), 28, fill="#FFFFFF", outline=f"#{LINE}", width=3)
    canvas.paste(laptop_fit, (lx, ly))
    px, py = 1370, 170
    frame_w, frame_h = 350, 830
    draw.rounded_rectangle((px, py, px + frame_w, py + frame_h), radius=54, fill="#17231F")
    screen = fit_image(phone, (318, 778), contain=False)
    screen = ImageOps.fit(screen, (318, 778), method=Image.Resampling.LANCZOS, centering=(0.5, 0.0))
    mask = Image.new("L", screen.size, 0)
    ImageDraw.Draw(mask).rounded_rectangle((0, 0, screen.width, screen.height), radius=38, fill=255)
    canvas.paste(screen, (px + 16, py + 26), mask)
    draw.rounded_rectangle((px + 128, py + 12, px + 222, py + 22), radius=6, fill="#3C4A45")
    draw.text((80, 980), "Laptop: rail trái + nội dung chính + hàng đợi ôn tập", font=pil_font(22, semi=True), fill=f"#{FOREST_DARK}")
    draw.text((760, 980), "Điện thoại: một cột + CTA lớn + thanh điều hướng dưới", font=pil_font(22, semi=True), fill=f"#{FOREST_DARK}")
    path = ASSETS / "responsive-overview.png"
    canvas.save(path, quality=94)
    return path


def crop_mobile_assets() -> None:
    for name in ["courses", "lesson", "practice", "vip", "admin"]:
        raw = Image.open(QA / f"{name}-mobile-raw.png").convert("RGB")
        raw.crop((0, 0, min(470, raw.width), min(1040, raw.height))).save(QA / f"{name}-mobile.png")


def build_screen_gallery(names: list[str], title: str, filename: str) -> Path:
    canvas = Image.new("RGB", (1800, 1120), f"#{PAPER}")
    draw = ImageDraw.Draw(canvas)
    draw.text((70, 45), title.upper(), font=pil_font(28, bold=True), fill=f"#{FOREST}")
    labels = {
        "courses": "THƯ VIỆN LỘ TRÌNH",
        "lesson": "KHÔNG GIAN BÀI HỌC",
        "practice": "ÔN TẬP THÔNG MINH",
        "vip": "GÓI VIP & CHÍNH SÁCH",
        "admin": "ADMIN VẬN HÀNH",
    }
    cell_w = 820
    x_positions = [70, 910]
    for idx, name in enumerate(names):
        x = x_positions[idx % 2]
        y = 120 + (idx // 2) * 500
        draw.text((x, y), labels[name], font=pil_font(22, semi=True), fill=f"#{FOREST_DARK}")
        lap = fit_image(Image.open(QA / f"{name}-laptop.png"), (620, 350))
        draw.rounded_rectangle((x, y + 48, x + 640, y + 418), radius=18, fill="#FFFFFF", outline=f"#{LINE}", width=2)
        canvas.paste(lap, (x + 10, y + 58))
        mob = fit_image(Image.open(QA / f"{name}-mobile.png"), (160, 350))
        mx = x + 650
        draw.rounded_rectangle((mx, y + 48, mx + 172, y + 418), radius=26, fill="#17231F")
        canvas.paste(mob, (mx + 6, y + 58))
    path = ASSETS / filename
    canvas.save(path, quality=94)
    return path


def build_roadmap() -> Path:
    canvas = Image.new("RGB", (1800, 1030), f"#{PAPER}")
    draw = ImageDraw.Draw(canvas)
    draw.text((70, 50), "LỘ TRÌNH 16 TUẦN TỪ PROTOTYPE ĐẾN PUBLIC BETA", font=pil_font(30, bold=True), fill=f"#{FOREST}")
    draw.text((70, 98), "Ưu tiên web trước; Flutter chỉ bắt đầu khi số liệu chứng minh người học quay lại.", font=pil_font(26, semi=True), fill=f"#{FOREST_DARK}")
    phases = [
        ("01", "Nền móng", "Tuần 1–2", "Auth · PostgreSQL · analytics", MINT),
        ("02", "Admin CMS", "Tuần 3–5", "Nháp · duyệt · xuất bản", "E8EEF5"),
        ("03", "Nội dung & audio", "Tuần 4–7", "3 lộ trình đầu · giọng chuẩn", GOLD_LIGHT),
        ("04", "Engine học", "Tuần 6–8", "Tiến độ · SRS · hàng đợi ôn", MINT),
        ("05", "VIP & SePay", "Tuần 8–10", "Đơn hàng · webhook · đối soát", "E8EEF5"),
        ("06", "Beta & tối ưu", "Tuần 11–16", "Bảo mật · retention · chuyển đổi", GOLD_LIGHT),
    ]
    x0, y0, gap, w, h = 70, 230, 18, 262, 430
    for i, (num, title, timing, detail, fill) in enumerate(phases):
        x = x0 + i * (w + gap)
        draw.rounded_rectangle((x, y0, x + w, y0 + h), radius=28, fill=f"#{fill}", outline=f"#{LINE}", width=3)
        draw.ellipse((x + 24, y0 + 24, x + 80, y0 + 80), fill=f"#{FOREST}")
        draw.text((x + 41, y0 + 38), num, font=pil_font(20, bold=True), fill="#FFFFFF", anchor="mm")
        draw_wrapped(draw, (x + 24, y0 + 110), title, pil_font(30, bold=True), f"#{FOREST_DARK}", w - 48, max_lines=2)
        draw.text((x + 24, y0 + 205), timing, font=pil_font(20, semi=True), fill=f"#{FOREST}")
        draw_wrapped(draw, (x + 24, y0 + 255), detail, pil_font(22), f"#{INK_SOFT}", w - 48, max_lines=4)
        if i < len(phases) - 1:
            arrow(draw, (x + w + 2, y0 + h // 2), (x + w + gap - 2, y0 + h // 2), FOREST, 4)
    gates = [
        "Gate A · Có dữ liệu thật thay mock",
        "Gate B · 30 bài đã được reviewer duyệt",
        "Gate C · Thanh toán idempotent, có đối soát",
        "Gate D · 4 tuần retention đủ tốt mới làm Flutter",
    ]
    y = 745
    for i, gate in enumerate(gates):
        gx = 70 + (i % 2) * 850
        gy = y + (i // 2) * 100
        draw.rounded_rectangle((gx, gy, gx + 790, gy + 70), radius=18, fill="#FFFFFF", outline=f"#{LINE}", width=2)
        draw.ellipse((gx + 20, gy + 20, gx + 50, gy + 50), fill=f"#{GOLD}")
        draw.text((gx + 68, gy + 21), gate, font=pil_font(22, semi=True), fill=f"#{FOREST_DARK}")
    path = ASSETS / "roadmap-16-weeks.png"
    canvas.save(path, quality=95)
    return path


def build_learning_bpmn() -> Path:
    canvas = Image.new("RGB", (1800, 1080), f"#{PAPER}")
    draw = ImageDraw.Draw(canvas)
    draw.text((55, 35), "BPMN RÚT GỌN · QUY TRÌNH HỌC", font=pil_font(30, bold=True), fill=f"#{FOREST}")
    lanes = [("NGƯỜI HỌC", 110, 370), ("HỆ THỐNG", 370, 680), ("DỮ LIỆU & ÔN TẬP", 680, 1000)]
    for label, top, bottom in lanes:
        draw.rectangle((40, top, 1760, bottom), fill="#FFFFFF", outline=f"#{LINE}", width=3)
        draw.rectangle((40, top, 190, bottom), fill=f"#{MINT_LIGHT}", outline=f"#{LINE}", width=3)
        draw_wrapped(draw, (68, top + 40), label, pil_font(21, bold=True), f"#{FOREST_DARK}", 95, max_lines=3)
    # learner lane
    draw.ellipse((225, 195, 275, 245), fill=f"#{FOREST}")
    rounded_task(draw, (320, 160, 545, 290), "Chọn lộ trình", "Theo mục tiêu công việc", accent=FOREST)
    rounded_task(draw, (625, 160, 865, 290), "Mở bài học", "Bài miễn phí hoặc VIP", accent=FOREST)
    rounded_task(draw, (990, 145, 1270, 305), "Học 3 phần", "Từ vựng · hội thoại · ghi chú", accent=FOREST)
    rounded_task(draw, (1385, 160, 1615, 290), "Đánh dấu xong", "Nhận bước tiếp theo", accent=FOREST)
    for a, b in [((275, 220), (320, 220)), ((545, 220), (625, 220)), ((865, 220), (990, 220)), ((1270, 220), (1385, 220))]:
        arrow(draw, a, b)
    # system lane
    rounded_task(draw, (600, 455, 850, 585), "Kiểm tra quyền", "free / VIP / khóa", fill="F9FBF9", accent=BLUE)
    diamond = [(940, 520), (1020, 445), (1100, 520), (1020, 595)]
    draw.polygon(diamond, fill=f"#{GOLD_LIGHT}", outline=f"#{GOLD}")
    draw_wrapped(draw, (970, 493), "Đủ quyền?", pil_font(22, bold=True), f"#{FOREST_DARK}", 105, max_lines=2)
    rounded_task(draw, (1160, 455, 1390, 585), "Ghi tiến độ", "Thời gian, phần đã xem", fill="F9FBF9", accent=BLUE)
    arrow(draw, (745, 370), (725, 455), BLUE)
    arrow(draw, (850, 520), (940, 520), BLUE)
    arrow(draw, (1100, 520), (1160, 520), BLUE)
    draw.text((1110, 485), "CÓ", font=pil_font(17, bold=True), fill=f"#{FOREST}")
    draw.text((970, 615), "KHÔNG → hiển thị lợi ích VIP, không mất vị trí học", font=pil_font(18, semi=True), fill=f"#{INK_SOFT}")
    arrow(draw, (1275, 455), (1500, 310), BLUE)
    # review lane
    rounded_task(draw, (390, 770, 665, 920), "Tạo review items", "Từ khó và từ vừa học", fill=MINT_LIGHT, accent=GOLD)
    rounded_task(draw, (780, 770, 1070, 920), "Xếp lịch SRS", "next_review_at + độ khó", fill=MINT_LIGHT, accent=GOLD)
    rounded_task(draw, (1200, 770, 1510, 920), "Phiên ôn kế tiếp", "Tự đánh giá rồi tăng/giảm khoảng cách", fill=MINT_LIGHT, accent=GOLD)
    arrow(draw, (1290, 585), (530, 770), GOLD)
    arrow(draw, (665, 845), (780, 845), GOLD)
    arrow(draw, (1070, 845), (1200, 845), GOLD)
    draw.ellipse((1580, 820, 1640, 880), outline=f"#{FOREST}", width=6)
    draw.ellipse((1590, 830, 1630, 870), fill=f"#{FOREST}")
    arrow(draw, (1510, 845), (1580, 850), FOREST)
    path = ASSETS / "learning-bpmn.png"
    canvas.save(path, quality=95)
    return path


def build_vip_bpmn() -> Path:
    canvas = Image.new("RGB", (1800, 1150), f"#{PAPER}")
    draw = ImageDraw.Draw(canvas)
    draw.text((55, 35), "BPMN RÚT GỌN · MUA VÀ KÍCH HOẠT VIP QUA SEPAY", font=pil_font(30, bold=True), fill=f"#{FOREST}")
    lane_defs = [("NGƯỜI HỌC", 105, 325), ("WEB / API", 325, 650), ("SEPAY / NGÂN HÀNG", 650, 885), ("ADMIN", 885, 1090)]
    for label, top, bottom in lane_defs:
        draw.rectangle((40, top, 1760, bottom), fill="#FFFFFF", outline=f"#{LINE}", width=3)
        draw.rectangle((40, top, 190, bottom), fill=f"#{MINT_LIGHT}", outline=f"#{LINE}", width=3)
        draw_wrapped(draw, (65, top + 36), label, pil_font(20, bold=True), f"#{FOREST_DARK}", 105, max_lines=3)
    # learner
    rounded_task(draw, (235, 145, 440, 275), "Đăng nhập", "Xác định người nhận VIP", accent=FOREST)
    rounded_task(draw, (520, 145, 735, 275), "Chọn gói", "1 / 6 / 12 tháng", accent=FOREST)
    rounded_task(draw, (835, 145, 1070, 275), "Quét QR", "Đúng số tiền + nội dung", accent=FOREST)
    rounded_task(draw, (1450, 145, 1660, 275), "Nhận VIP", "Có ngày hết hạn rõ ràng", accent=FOREST)
    arrow(draw, (440, 210), (520, 210)); arrow(draw, (735, 210), (835, 210))
    # web
    rounded_task(draw, (490, 410, 760, 555), "Tạo payment_order", "reference_code duy nhất · hết hạn", fill="F9FBF9", accent=BLUE)
    rounded_task(draw, (840, 410, 1110, 555), "Hiển thị QR", "Không đánh dấu đã trả ở client", fill="F9FBF9", accent=BLUE)
    diamond = [(1260, 485), (1340, 410), (1420, 485), (1340, 560)]
    draw.polygon(diamond, fill=f"#{GOLD_LIGHT}", outline=f"#{GOLD}")
    draw_wrapped(draw, (1285, 457), "Webhook hợp lệ?", pil_font(20, bold=True), f"#{FOREST_DARK}", 110, max_lines=2)
    rounded_task(draw, (1480, 405, 1690, 565), "Kích hoạt", "Transaction id duy nhất · atomic", fill=MINT_LIGHT, accent=FOREST)
    arrow(draw, (625, 325), (625, 410), BLUE); arrow(draw, (760, 485), (840, 485), BLUE); arrow(draw, (1110, 485), (1260, 485), BLUE); arrow(draw, (1420, 485), (1480, 485), BLUE); arrow(draw, (1585, 405), (1555, 275), FOREST)
    # provider
    rounded_task(draw, (855, 700, 1115, 825), "Ghi nhận chuyển khoản", "Giao dịch ngân hàng thực", fill=GOLD_LIGHT, accent=GOLD)
    rounded_task(draw, (1210, 700, 1465, 825), "Gửi webhook", "Payload + chữ ký / token", fill=GOLD_LIGHT, accent=GOLD)
    arrow(draw, (955, 275), (980, 700), GOLD); arrow(draw, (1115, 760), (1210, 760), GOLD); arrow(draw, (1335, 700), (1340, 560), GOLD)
    # admin exception
    rounded_task(draw, (1030, 930, 1335, 1045), "Hàng đợi ngoại lệ", "Sai tiền · sai nội dung · webhook lỗi", fill="FFF8E9", accent=GOLD)
    rounded_task(draw, (1430, 930, 1685, 1045), "Đối soát thủ công", "Có audit log và người xử lý", fill="FFF8E9", accent=GOLD)
    arrow(draw, (1340, 560), (1180, 930), GOLD)
    draw.text((1220, 600), "KHÔNG", font=pil_font(17, bold=True), fill=f"#{GOLD}")
    arrow(draw, (1335, 988), (1430, 988), GOLD)
    path = ASSETS / "vip-bpmn.png"
    canvas.save(path, quality=95)
    return path


def build_content_bpmn() -> Path:
    canvas = Image.new("RGB", (1800, 1010), f"#{PAPER}")
    draw = ImageDraw.Draw(canvas)
    draw.text((55, 35), "BPMN RÚT GỌN · BIÊN SOẠN VÀ XUẤT BẢN NỘI DUNG", font=pil_font(30, bold=True), fill=f"#{FOREST}")
    lanes = [("EDITOR", 110, 390), ("REVIEWER TIẾNG TRUNG", 390, 680), ("HỆ THỐNG / ADMIN", 680, 950)]
    for label, top, bottom in lanes:
        draw.rectangle((40, top, 1760, bottom), fill="#FFFFFF", outline=f"#{LINE}", width=3)
        draw.rectangle((40, top, 190, bottom), fill=f"#{MINT_LIGHT}", outline=f"#{LINE}", width=3)
        draw_wrapped(draw, (62, top + 40), label, pil_font(20, bold=True), f"#{FOREST_DARK}", 110, max_lines=4)
    rounded_task(draw, (240, 175, 485, 325), "Soạn bài", "Tình huống · từ · hội thoại", accent=FOREST)
    rounded_task(draw, (590, 175, 845, 325), "Tự kiểm tra", "Pinyin · nghĩa · nguồn", accent=FOREST)
    rounded_task(draw, (1260, 175, 1515, 325), "Sửa theo góp ý", "Ghi change note", accent=FOREST)
    arrow(draw, (485, 250), (590, 250)); arrow(draw, (845, 250), (960, 465), BLUE)
    rounded_task(draw, (920, 455, 1185, 600), "Duyệt chuyên môn", "Tự nhiên · đúng ngữ cảnh", fill="F9FBF9", accent=BLUE)
    diamond = [(1275, 525), (1350, 455), (1425, 525), (1350, 595)]
    draw.polygon(diamond, fill=f"#{GOLD_LIGHT}", outline=f"#{GOLD}")
    draw_wrapped(draw, (1308, 497), "Đạt?", pil_font(22, bold=True), f"#{FOREST_DARK}", 85, max_lines=2)
    arrow(draw, (1185, 525), (1275, 525), BLUE)
    arrow(draw, (1350, 455), (1380, 325), GOLD)
    draw.text((1370, 415), "CHƯA", font=pil_font(17, bold=True), fill=f"#{GOLD}")
    rounded_task(draw, (650, 745, 930, 880), "Tạo bản xem trước", "Không lộ URL công khai", fill=MINT_LIGHT, accent=FOREST)
    rounded_task(draw, (1035, 745, 1300, 880), "Xuất bản version", "Snapshot + người xuất bản", fill=MINT_LIGHT, accent=FOREST)
    rounded_task(draw, (1405, 745, 1645, 880), "Ghi audit log", "Có thể rollback", fill=MINT_LIGHT, accent=FOREST)
    arrow(draw, (1350, 595), (790, 745), FOREST); draw.text((1185, 625), "ĐẠT", font=pil_font(17, bold=True), fill=f"#{FOREST}")
    arrow(draw, (930, 812), (1035, 812)); arrow(draw, (1300, 812), (1405, 812))
    path = ASSETS / "content-bpmn.png"
    canvas.save(path, quality=95)
    return path


def build_architecture() -> Path:
    canvas = Image.new("RGB", (1800, 1020), f"#{PAPER}")
    draw = ImageDraw.Draw(canvas)
    draw.text((70, 45), "KIẾN TRÚC ĐỀ XUẤT · WEB TRƯỚC, DÙNG CHUNG API CHO FLUTTER SAU", font=pil_font(30, bold=True), fill=f"#{FOREST}")
    layers = [
        ("TRẢI NGHIỆM", [("Next.js Web", "Responsive · SEO · PWA"), ("Flutter (sau gate)", "Android/iOS dùng chung API")], MINT_LIGHT),
        ("ỨNG DỤNG", [("App Router", "RSC · Server Actions"), ("Route Handlers", "Auth · progress · admin"), ("Payment module", "Order · webhook · idempotency")], "E8EEF5"),
        ("DỮ LIỆU", [("PostgreSQL", "Drizzle ORM · migration"), ("Object storage", "Audio · ảnh · export"), ("Job queue", "Webhook retry · email · audio")], GOLD_LIGHT),
        ("VẬN HÀNH", [("Analytics", "activation · retention"), ("Observability", "logs · alerts · audit"), ("Backup", "PITR · restore drill")], MINT_LIGHT),
    ]
    top = 135
    for li, (label, items, fill) in enumerate(layers):
        y = top + li * 215
        draw.rounded_rectangle((70, y, 1730, y + 180), radius=26, fill=f"#{fill}", outline=f"#{LINE}", width=3)
        draw.text((100, y + 28), label, font=pil_font(22, bold=True), fill=f"#{FOREST}")
        item_w = 445
        start_x = 390
        for ii, (title, subtitle) in enumerate(items):
            x = start_x + ii * (item_w + 35)
            rounded_task(draw, (x, y + 25, x + item_w, y + 155), title, subtitle, fill=WHITE, accent=FOREST if li == 0 else BLUE if li == 1 else GOLD)
        if li < len(layers) - 1:
            arrow(draw, (900, y + 180), (900, y + 215), FOREST, 4)
    path = ASSETS / "architecture.png"
    canvas.save(path, quality=95)
    return path


def build_data_map() -> Path:
    canvas = Image.new("RGB", (1800, 1120), f"#{PAPER}")
    draw = ImageDraw.Draw(canvas)
    draw.text((65, 40), "POSTGRESQL · BẢN ĐỒ MIỀN DỮ LIỆU", font=pil_font(30, bold=True), fill=f"#{FOREST}")
    groups = [
        ("IDENTITY", ["users", "audit_logs"], 80, 130, MINT_LIGHT),
        ("CONTENT", ["courses", "modules", "lessons", "vocabulary", "lesson_vocabulary", "content_versions"], 450, 130, "E8EEF5"),
        ("LEARNING", ["lesson_progress", "review_items"], 80, 620, GOLD_LIGHT),
        ("COMMERCE", ["vip_plans", "subscriptions", "payment_orders", "payment_events"], 930, 620, MINT_LIGHT),
    ]
    boxes: dict[str, tuple[int, int, int, int]] = {}
    for label, tables, x, y, fill in groups:
        cols = 3 if label == "CONTENT" else 2
        rows = (len(tables) + cols - 1) // cols
        group_w = 790 if label in {"CONTENT", "COMMERCE"} else 620
        group_h = 390 if label == "CONTENT" else 360
        draw.rounded_rectangle((x, y, x + group_w, y + group_h), radius=28, fill=f"#{fill}", outline=f"#{LINE}", width=3)
        draw.text((x + 24, y + 20), label, font=pil_font(21, bold=True), fill=f"#{FOREST}")
        for i, table in enumerate(tables):
            c, r = i % cols, i // cols
            tw = (group_w - 60 - (cols - 1) * 16) // cols
            tx = x + 24 + c * (tw + 16)
            ty = y + 70 + r * 105
            box = (tx, ty, tx + tw, ty + 78)
            draw.rounded_rectangle(box, radius=16, fill="#FFFFFF", outline=f"#{LINE}", width=2)
            draw.text((tx + 15, ty + 24), table, font=pil_font(20, semi=True), fill=f"#{FOREST_DARK}")
            boxes[table] = box
    relations = [
        ("users", "lesson_progress"), ("users", "subscriptions"), ("courses", "modules"),
        ("modules", "lessons"), ("lessons", "lesson_vocabulary"), ("vocabulary", "lesson_vocabulary"),
        ("lessons", "content_versions"), ("vocabulary", "review_items"), ("vip_plans", "subscriptions"),
        ("vip_plans", "payment_orders"), ("payment_orders", "payment_events"),
    ]
    for src, dst in relations:
        if src not in boxes or dst not in boxes:
            continue
        a, b = boxes[src], boxes[dst]
        start = ((a[0] + a[2]) // 2, a[3])
        end = ((b[0] + b[2]) // 2, b[1])
        draw.line([start, end], fill=f"#{FOREST}", width=3)
    draw.text((80, 1035), "Khóa quan trọng: email, slug, reference_code, provider_transaction_id, provider_event_id và (user_id, lesson_id).", font=pil_font(22, semi=True), fill=f"#{FOREST_DARK}")
    path = ASSETS / "data-map.png"
    canvas.save(path, quality=95)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width: int = 9360) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width))
    tbl_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(INK_SOFT)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(FOREST_DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for style_name, size, color, before, after in [
        ("Title", 34, FOREST_DARK, 0, 12),
        ("Subtitle", 15, INK_SOFT, 0, 8),
        ("Heading 1", 18, BLUE, 18, 9),
        ("Heading 2", 14, BLUE, 14, 7),
        ("Heading 3", 11.5, FOREST_DARK, 10, 5),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = style_name != "Subtitle"
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        if style_name == "Heading 1":
            st.paragraph_format.page_break_before = True

    header = section.header
    p = header.paragraphs[0]
    p.text = "HANZIWORK  ·  PRODUCT BLUEPRINT  ·  v0.2"
    p.style = styles["Normal"]
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(FOREST)
    p.paragraph_format.space_after = Pt(0)
    footer = section.footer
    add_page_field(footer.paragraphs[0])


def add_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text.upper())
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(FOREST)


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.add_run(item)


def add_callout(doc: Document, title: str, text: str, fill: str = MINT_LIGHT, icon: str = "QUYẾT ĐỊNH") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{icon}  ·  {title}")
    run.font.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(FOREST)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None,
              header_fill: str = "E8EEF5"):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table)
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_repeat_no_split(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(FOREST_DARK)
        if widths:
            cell.width = Inches(widths[i] / 1440)
    for row in rows:
        cells = table.add_row().cells
        set_repeat_no_split(table.rows[-1])
        for i, text in enumerate(row):
            cell = cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(str(text))
            r.font.size = Pt(8.8)
            if widths:
                cell.width = Inches(widths[i] / 1440)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_picture(doc: Document, path: Path, width: float = 6.65, caption: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = bool(caption)
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        cr = cp.runs[0]
        cr.italic = True
        cr.font.size = Pt(8.5)
        cr.font.color.rgb = RGBColor.from_string(INK_SOFT)


def page_break(doc: Document) -> None:
    # Major sections use Heading 1 with page_break_before. Keeping this helper as
    # a semantic marker avoids blank pages when a preceding table fills a page.
    return None


def build_document() -> Path:
    crop_mobile_assets()
    responsive = build_responsive_overview()
    learning_gallery = build_screen_gallery(["courses", "lesson", "practice"], "Các mặt học tập", "gallery-learning.png")
    ops_gallery = build_screen_gallery(["vip", "admin"], "Thương mại và vận hành", "gallery-ops.png")
    roadmap = build_roadmap()
    learning_bpmn = build_learning_bpmn()
    vip_bpmn = build_vip_bpmn()
    content_bpmn = build_content_bpmn()
    architecture = build_architecture()
    data_map = build_data_map()

    doc = Document()
    configure_document(doc)

    # Cover
    add_label(doc, "Product blueprint · UX · BPMN · PostgreSQL · VIP")
    p = doc.add_paragraph(style="Title")
    p.add_run("HanziWork")
    p2 = doc.add_paragraph(style="Subtitle")
    p2.add_run("Lộ trình sản phẩm, quy trình nghiệp vụ và đặc tả giao diện responsive")
    doc.add_paragraph()
    add_callout(
        doc,
        "Tuyên bố sản phẩm",
        "Tiếng Trung chuyên ngành theo tình huống thực tế cho người đang đi làm — học đúng việc cần làm, ôn đúng từ còn yếu và mở rộng sang ngành mới mà không phải học lan man.",
        fill=MINT,
        icon="TẦM NHÌN",
    )
    add_picture(doc, responsive, width=6.65)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("Phiên bản 0.2  ·  31/07/2026  ·  Chủ dự án: Gia Huy & cộng sự").font.size = Pt(9)
    p3.runs[0].font.color.rgb = RGBColor.from_string(INK_SOFT)

    page_break(doc)
    doc.add_heading("Tóm tắt điều hành", level=1)
    add_body(doc, "Prototype hiện đã có sáu tuyến chính: trang học hôm nay, thư viện lộ trình, không gian bài học, ôn tập, VIP và Admin. Phương án thiết kế số 3 đã được triển khai thành một trạm học tập responsive: laptop dùng rail trái và ba vùng ưu tiên; điện thoại chuyển sang một cột với nút học lớn và điều hướng dưới.")
    add_table(doc, ["Hiện trạng", "Quyết định", "Kết quả cần đạt"], [
        ["UI mẫu dùng dữ liệu tĩnh", "Dùng prototype để khóa luồng trước khi nối dữ liệu thật", "Giảm sửa lại giao diện khi backend bắt đầu"],
        ["Chưa có audio", "Không giả lập phát âm; hiển thị trạng thái sắp có", "Minh bạch và tránh trải nghiệm sai"],
        ["Đội 2 người", "Next.js full-stack + PostgreSQL trong MVP", "Một codebase, vận hành gọn"],
        ["Flutter để sau", "Chỉ khởi động khi web có retention và VIP thật", "Không chia nguồn lực quá sớm"],
    ], widths=[2900, 3100, 3360])
    doc.add_heading("Ba việc cần làm ngay", level=2)
    add_numbered(doc, [
        "Nối đăng nhập, PostgreSQL và quyền truy cập để thay dữ liệu minh họa.",
        "Xây Admin CMS có trạng thái nháp → duyệt → xem trước → xuất bản; đồng thời làm 30 bài đầu cho ba lộ trình dễ kiểm chứng.",
        "Sau khi luồng học ổn định mới tích hợp đơn hàng, QR SePay, webhook idempotent và chính sách bán VIP.",
    ])
    add_callout(doc, "Nguyên tắc cốt lõi", "Web thắng bằng nội dung đúng và nhịp học quay lại; không thắng nhờ có thật nhiều tính năng. Audio chuẩn và reviewer tiếng Trung quan trọng hơn chấm điểm nói bằng AI ở giai đoạn đầu.", fill=GOLD_LIGHT, icon="ƯU TIÊN")

    doc.add_heading("Mục lục", level=2)
    add_table(doc, ["Phần", "Nội dung"], [
        ["01", "Sản phẩm, phạm vi MVP và hệ giao diện"],
        ["02", "Nghiệp vụ từng trang và BPMN học tập"],
        ["03", "VIP, SePay, chính sách và quản trị nội dung"],
        ["04", "Lộ trình 16 tuần, kiến trúc và PostgreSQL"],
        ["05", "Nội dung, đo lường, Flutter gate và checklist ra mắt"],
    ], widths=[1200, 8160])

    page_break(doc)
    doc.add_heading("01 · Sản phẩm và phạm vi MVP", level=1)
    doc.add_heading("Đối tượng và công việc cần hoàn thành", level=2)
    add_table(doc, ["Nhóm người học", "Nhu cầu thực tế", "Giá trị HanziWork"], [
        ["Nhân viên mới làm với đối tác Trung Quốc", "Chào hỏi, giao việc, hỏi tiến độ, xác nhận", "Mẫu câu dùng ngay trong công việc"],
        ["Người chuyển ngành", "Nắm vốn từ của môi trường mới", "Lộ trình theo tình huống thay vì giáo trình tổng quát"],
        ["Nhân viên nhà máy / kho vận", "Hiểu thuật ngữ thao tác và báo cáo", "Bài ngắn 8–12 phút, ôn từ khó"],
        ["Nhân viên bán hàng / CSKH", "Tư vấn, phản hồi, đổi trả, lịch hẹn", "Hội thoại hai chiều và ghi chú sắc thái"],
    ], widths=[2700, 3200, 3460])
    doc.add_heading("Phạm vi V1", level=2)
    add_table(doc, ["Có trong V1", "Chưa làm trong V1"], [
        ["Web responsive; email/Google login; hồ sơ học", "Chấm điểm phát âm bằng AI"],
        ["3 lộ trình đầu, mỗi lộ trình 10–12 bài", "Mạng xã hội, bảng xếp hạng công khai"],
        ["Từ vựng, hội thoại, ghi chú, audio chuẩn", "Livestream hoặc lớp học trực tiếp"],
        ["Tiến độ, flashcard, SRS cơ bản", "Offline đầy đủ và đồng bộ phức tạp"],
        ["VIP 1/6/12 tháng, SePay, đối soát", "Tự động gia hạn thẻ"],
        ["Admin CMS, reviewer, version, audit", "Flutter trước khi đạt product gate"],
    ], widths=[4680, 4680])
    add_callout(doc, "Ba lộ trình nên làm đầu", "Văn phòng & hành chính; nhà máy & sản xuất; kho vận & logistics. Đây là các ngữ cảnh có nhiều tình huống lặp lại, dễ kiểm chứng với người thật và ít rủi ro hơn y tế, pháp lý hay kế toán chuyên sâu.")

    page_break(doc)
    doc.add_heading("Hệ giao diện đã chọn", level=1)
    add_picture(doc, responsive, width=5.55, caption="Phương án 3 đã triển khai: laptop và điện thoại dùng cùng thứ tự ưu tiên, khác cách sắp xếp.")
    doc.add_heading("Ngôn ngữ thiết kế", level=2)
    add_table(doc, ["Thành phần", "Quy tắc"], [
        ["Màu", "Pine #0F6A58 cho hành động; cream #F7F6F1 làm nền; mint cho trạng thái học; gold chỉ dùng cho VIP/cảnh báo."],
        ["Chữ", "Sans dễ đọc cho tiếng Việt; serif cho Hán tự; cỡ chữ nội dung tối thiểu khoảng 14–16 px trên mobile."],
        ["Bố cục", "Laptop: rail 88 px + nội dung chính + cột phụ. Mobile: một cột, CTA rộng, bottom nav 4 mục."],
        ["Thẻ", "Bo góc vừa, viền mảnh, ít bóng; một màn hình chỉ có một hành động chính nổi bật."],
        ["Trạng thái", "Không giả lập tính năng chưa có. Audio chưa sẵn sàng phải disabled và có nhãn giải thích."],
    ], widths=[1800, 7560])
    doc.add_heading("Quy tắc responsive", level=2)
    add_table(doc, ["Breakpoint", "Hành vi"], [
        ["> 1120 px", "Rail trái cố định; bài học ở giữa; ôn tập và mục tiêu ở cột phải."],
        ["921–1120 px", "Cột phải chuyển xuống dưới; khu học chính giữ ưu tiên."],
        ["721–920 px", "Ẩn rail; topbar gọn; nội dung một cột rộng."],
        ["≤ 720 px", "Bottom nav cố định; thẻ bài học xếp dọc; CTA cao tối thiểu 52 px; tránh bảng tràn."],
        ["≤ 390 px", "Giảm padding nhưng không giảm vùng chạm; tiêu đề co bằng clamp; text không bị cắt."],
    ], widths=[1900, 7460])

    page_break(doc)
    doc.add_heading("02 · Nghiệp vụ từng trang", level=1)
    add_picture(doc, learning_gallery, width=5.0, caption="Các mặt học tập đã được kiểm tra ở viewport laptop và điện thoại.")
    doc.add_heading("Trang Hôm nay", level=2)
    add_table(doc, ["Mục", "Đặc tả"], [
        ["Mục tiêu", "Đưa người học vào bài tiếp theo trong dưới 10 giây."],
        ["Khối chính", "Lời chào; bài hôm nay; từ trọng tâm; CTA bắt đầu; ôn đến hạn; mục tiêu tuần; gợi ý VIP."],
        ["Luồng", "Đăng nhập → hệ thống lấy bài tiếp theo + review due → người học chọn học hoặc ôn → ghi sự kiện."],
        ["Trạng thái rỗng", "Chưa chọn ngành: đề nghị chọn lộ trình. Không có review: gợi ý bài mới. Hết VIP: giữ tiến độ, khóa phần VIP."],
        ["Dữ liệu", "users, lesson_progress, review_items, subscriptions."],
    ], widths=[1800, 7560])
    doc.add_heading("Trang Lộ trình", level=2)
    add_table(doc, ["Mục", "Đặc tả"], [
        ["Mục tiêu", "Giúp chọn đúng ngành theo công việc cần làm, không theo cấp độ ngữ pháp."],
        ["Khối chính", "Search; chip ngành; thẻ lộ trình; số bài/thời lượng; số bài học thử; trạng thái đã học."],
        ["Luồng", "Nhập từ khóa/lọc → xem thẻ → mở lộ trình → xem bài miễn phí hoặc kiểm tra VIP."],
        ["Trạng thái", "Không có kết quả; lộ trình nháp không hiển thị; thẻ VIP phải nói rõ số bài miễn phí."],
        ["Dữ liệu", "courses, modules, lessons, lesson_progress."],
    ], widths=[1800, 7560])

    page_break(doc)
    doc.add_heading("Trang Bài học", level=1)
    add_table(doc, ["Mục", "Đặc tả"], [
        ["Mục tiêu", "Hoàn thành một tình huống công việc trong 8–12 phút."],
        ["Khối chính", "Danh sách bài; từ vựng; ví dụ; hội thoại; ghi chú dùng từ; audio; hoàn thành; bài tiếp theo."],
        ["Luồng", "Kiểm tra quyền → tải nội dung published → ghi last_opened_at → học các phần → hoàn thành → tạo review items."],
        ["Quy tắc", "Bài miễn phí mở cho mọi tài khoản. Bài VIP không mất vị trí khi bị khóa. Audio lỗi phải có retry và transcript."],
        ["Dữ liệu", "lessons.content, vocabulary, lesson_vocabulary, lesson_progress, review_items."],
    ], widths=[1800, 7560])
    doc.add_heading("Trang Ôn tập", level=2)
    add_table(doc, ["Mục", "Đặc tả"], [
        ["Mục tiêu", "Ôn đúng từ đến hạn; người học tự đánh giá sau khi mở đáp án."],
        ["Khối chính", "Số từ đến hạn; tiến độ phiên; thẻ Hán tự/pinyin/nghĩa; cần ôn lại/đã nhớ; kết quả."],
        ["Luồng", "Lấy review_items next_review_at ≤ now → lật thẻ → đánh giá → cập nhật interval/ease → lấy từ tiếp theo."],
        ["Quy tắc", "Không cho đánh giá trước khi lật thẻ. Sai: lặp lại sớm. Đúng liên tiếp: tăng khoảng cách. Giới hạn phiên 10–20 từ."],
        ["Dữ liệu", "review_items, vocabulary; event practice_answered."],
    ], widths=[1800, 7560])
    add_picture(doc, learning_bpmn, width=5.95, caption="BPMN rút gọn của hành trình học và vòng lặp ôn tập.")

    page_break(doc)
    doc.add_heading("Trang VIP và Trang Admin", level=1)
    add_picture(doc, ops_gallery, width=5.65, caption="Hai mặt vận hành: chuyển đổi trả phí và quản trị nội dung/dữ liệu.")
    doc.add_heading("Trang VIP", level=2)
    add_table(doc, ["Mục", "Đặc tả"], [
        ["Mục tiêu", "So sánh gói minh bạch, tạo đơn đúng người và kích hoạt tự động sau thanh toán."],
        ["Khối chính", "Quyền lợi; 3 gói; thời hạn/giá; chính sách; checkout; QR; trạng thái đơn; hỗ trợ."],
        ["Trạng thái", "Chưa đăng nhập; đơn chờ; đã thanh toán; webhook chậm; sai số tiền; đơn hết hạn; đã có VIP."],
        ["Dữ liệu", "vip_plans, payment_orders, payment_events, subscriptions, audit_logs."],
    ], widths=[1800, 7560])
    doc.add_heading("Trang Admin", level=2)
    add_table(doc, ["Mục", "Đặc tả"], [
        ["Mục tiêu", "Cho đội hai người vận hành nội dung, người học và giao dịch mà không sửa trực tiếp DB."],
        ["Khối chính", "Dashboard; CMS lộ trình/bài/từ; review queue; audio; người dùng; đơn hàng; audit; cấu hình gói."],
        ["Vai trò", "Editor soạn; Reviewer duyệt tiếng Trung; Admin xuất bản, xử lý thanh toán và phân quyền."],
        ["Bảo vệ", "RBAC server-side; re-auth cho tác vụ nhạy cảm; audit log; không hiển thị password hash hay webhook secret."],
    ], widths=[1800, 7560])

    page_break(doc)
    doc.add_heading("03 · Chính sách VIP cụ thể cho MVP", level=1)
    add_table(doc, ["Quyền lợi", "Miễn phí", "VIP"], [
        ["Bài học", "3–5 bài mẫu mỗi lộ trình", "Toàn bộ bài đã xuất bản trong mọi lộ trình"],
        ["Từ vựng & ví dụ", "Đầy đủ trong bài mẫu", "Đầy đủ trong tất cả bài"],
        ["Hội thoại & ghi chú", "Một phần giới thiệu", "Toàn bộ nội dung"],
        ["Audio", "Audio của bài mẫu khi đã sản xuất", "Toàn bộ audio sẵn có; không cam kết AI speech scoring"],
        ["Ôn tập", "Phiên giới hạn hằng ngày", "Không giới hạn hợp lý + SRS đầy đủ"],
        ["Tiến độ", "Lưu bài mẫu", "Lưu toàn bộ và đồng bộ thiết bị"],
        ["Nội dung mới", "Xem mô tả", "Được mở trong thời hạn gói nếu thuộc thư viện VIP"],
    ], widths=[2500, 3100, 3760])
    doc.add_heading("Gói và cách gia hạn", level=2)
    add_table(doc, ["Gói mẫu", "Giá prototype", "Cách vận hành MVP"], [
        ["1 tháng", "79.000đ", "Thanh toán một lần; không tự động gia hạn"],
        ["6 tháng", "329.000đ", "Thanh toán một lần; cộng nối tiếp nếu mua khi còn hạn"],
        ["12 tháng", "549.000đ", "Thanh toán một lần; giá cần A/B test trước khi công bố"],
    ], widths=[2100, 2300, 4960])
    add_callout(doc, "Giá và chính sách đang là đề xuất sản phẩm", "Trước khi nhận tiền thật cần xác nhận chủ thể kinh doanh, tài khoản nhận tiền, nội dung điều khoản, bảo mật, hoàn tiền và thủ tục thương mại điện tử phù hợp. Tài liệu này không thay thế tư vấn thuế hoặc pháp lý.", fill=GOLD_LIGHT, icon="LƯU Ý")
    doc.add_heading("Chính sách vận hành đề xuất", level=2)
    add_bullets(doc, [
        "Kích hoạt: chỉ sau khi backend xác minh webhook, đúng reference code, đúng số tiền và transaction id chưa xử lý.",
        "Gia hạn: cộng ngày vào ends_at nếu gói hiện tại còn active; không làm ngắn thời hạn đang có.",
        "Hoàn tiền/điều chỉnh: hỗ trợ giao dịch trùng, bị trừ tiền nhưng không kích hoạt, hoặc lỗi hệ thống chưa thể khắc phục; mọi xử lý có audit log.",
        "Tài khoản: một gói gắn với một tài khoản; giới hạn chia sẻ nên áp dụng mềm ở MVP, cảnh báo trước khi khóa.",
        "Đơn hết hạn: QR hết hạn sau thời gian cấu hình; giao dịch đến muộn chuyển manual_review, không tự mất tiền của người học.",
    ])
    add_picture(doc, vip_bpmn, width=6.2, caption="Luồng VIP tách rõ client, backend, nhà cung cấp thanh toán và xử lý ngoại lệ.")

    page_break(doc)
    doc.add_heading("Quy trình Admin và chất lượng nội dung", level=1)
    add_picture(doc, content_bpmn, width=6.05, caption="Mỗi lần xuất bản tạo version và audit log; nội dung chưa duyệt không được xuất hiện ở app.")
    doc.add_heading("Definition of Done cho một bài học", level=2)
    add_table(doc, ["Cổng chất lượng", "Điều kiện đạt"], [
        ["Tình huống", "Có người nói, bối cảnh, mục tiêu giao tiếp và kết quả mong đợi."],
        ["Từ vựng", "6–10 từ; Hán tự, pinyin, nghĩa Việt, ví dụ Trung–Việt; không trùng vô ích."],
        ["Hội thoại", "4–8 lượt; tự nhiên; đúng quan hệ vai vế; có bản dịch."],
        ["Ghi chú", "Nêu sắc thái, lỗi hay gặp hoặc khác biệt trang trọng/thân mật."],
        ["Review", "Người kiểm tra tiếng Trung duyệt; editor sửa hết góp ý blocking."],
        ["Kỹ thuật", "Preview đúng mobile/laptop; audio URL hợp lệ; status published; version snapshot có đủ."],
    ], widths=[2200, 7160])
    doc.add_heading("Phân công đội hai người", level=2)
    add_table(doc, ["Vai trò", "Trách nhiệm chính", "Không nên tự duyệt"], [
        ["Bạn A · Product/Frontend", "UX, Next.js UI, nội dung khung, analytics, phỏng vấn người học", "Chất lượng tiếng Trung của chính bài mình soạn"],
        ["Bạn B · Backend/Ops", "DB, auth, Admin, payment, deploy, backup, support", "Giao dịch ngoại lệ do chính mình sửa không có audit"],
        ["Reviewer cộng tác", "Ngữ cảnh, pinyin, sắc thái, hội thoại, audio script", "Xuất bản hoặc chỉnh giá"],
    ], widths=[2200, 4700, 2460])

    page_break(doc)
    doc.add_heading("04 · Lộ trình 16 tuần", level=1)
    add_picture(doc, roadmap, caption="Kế hoạch có phần chồng lấn để hai người làm song song, nhưng mỗi giai đoạn đều có cổng nghiệm thu.")
    add_table(doc, ["Giai đoạn", "Đầu ra bắt buộc", "Điều kiện qua cổng"], [
        ["Tuần 1–2 · Nền móng", "Auth; PostgreSQL; migration; seed; analytics events", "Tài khoản thật học được bài mẫu và lưu tiến độ"],
        ["Tuần 3–5 · Admin CMS", "CRUD; draft/review/published; RBAC; preview", "Không sửa DB tay để xuất bản bài"],
        ["Tuần 4–7 · Nội dung", "30 bài đầu; 180–300 từ; audio batch 1", "100% bài có reviewer và checklist"],
        ["Tuần 6–8 · Learning", "Progress; review_items; SRS; dashboard thật", "Resume đúng và ôn lại không mất dữ liệu"],
        ["Tuần 8–10 · VIP", "Plans; orders; QR; webhook; audit; policy pages", "Retry webhook không kích hoạt hai lần"],
        ["Tuần 11–12 · Closed beta", "20–50 người dùng; support loop; fix P0/P1", "≥60% hoàn thành bài đầu; payment sandbox ổn"],
        ["Tuần 13–16 · Public beta", "SEO; performance; backup; báo cáo tuần", "Có số liệu retention và conversion để quyết định"],
    ], widths=[2400, 4300, 2660])

    page_break(doc)
    doc.add_heading("Kiến trúc triển khai", level=1)
    add_picture(doc, architecture, width=5.8, caption="Trong MVP có thể dùng Next.js cho cả frontend và backend. Flutter sau này gọi cùng Route Handlers/API.")
    doc.add_heading("Quyết định kỹ thuật", level=2)
    add_table(doc, ["Hạng mục", "Khuyến nghị", "Lý do"], [
        ["Frontend", "Next.js App Router + TypeScript", "SSR/RSC, responsive, một codebase"],
        ["Backend MVP", "Server Actions + Route Handlers trong Next.js", "Đủ cho đội nhỏ; dễ chia module commerce/learning/admin"],
        ["Database", "PostgreSQL + Drizzle", "Quan hệ rõ, migration kiểm soát, phù hợp payment/audit"],
        ["Audio", "Object storage + CDN; lưu URL trong vocabulary", "Không nhét binary vào PostgreSQL"],
        ["Auth", "Provider uy tín hoặc thư viện có session server-side", "Không tự viết crypto/password flow"],
        ["Payment", "Module riêng; webhook xác thực; idempotency; transaction DB", "Tránh kích hoạt trùng và sai người"],
        ["Flutter", "Chỉ gọi API; không truy cập DB trực tiếp", "Giữ quy tắc nghiệp vụ ở server"],
    ], widths=[1900, 3300, 4160])
    add_callout(doc, "Khi nào tách backend riêng?", "Chỉ tách NestJS/Fastify khi job nền, realtime, nhiều client hoặc đội backend đã lớn đến mức module trong Next.js gây cản trở. Với hai người, tách sớm sẽ tăng chi phí triển khai và debug.")

    page_break(doc)
    doc.add_heading("PostgreSQL và mô hình dữ liệu", level=1)
    add_picture(doc, data_map, width=5.7, caption="Schema hiện đã bao phủ identity, content, learning, commerce và audit.")
    add_table(doc, ["Miền", "Bảng", "Ghi chú quan trọng"], [
        ["Identity", "users, audit_logs", "Email unique; role server-side; is_active; actor/action/entity metadata"],
        ["Content", "courses, modules, lessons", "Slug unique; sort_order; status; published_at; JSONB content có version"],
        ["Vocabulary", "vocabulary, lesson_vocabulary", "Tái sử dụng từ; audio_url; tags; thứ tự trong bài"],
        ["Learning", "lesson_progress, review_items", "Composite key theo user; index review due"],
        ["Commerce", "vip_plans, subscriptions", "Không hard-code giá ở client; starts_at/ends_at/status"],
        ["Payment", "payment_orders, payment_events", "reference và provider transaction unique; giữ payload thô"],
        ["Governance", "content_versions, audit_logs", "Snapshot để rollback; ai làm gì, lúc nào"],
    ], widths=[1700, 3000, 4660])
    doc.add_heading("Ràng buộc và bảo mật bắt buộc", level=2)
    add_bullets(doc, [
        "Webhook xử lý trong transaction: ghi event → khóa order → kiểm tra amount/status → activate subscription → commit.",
        "Không tin plan_id, amount_vnd, role hoặc trạng thái VIP gửi từ client; luôn lấy lại từ database.",
        "Mỗi migration được review và backup trước thay đổi phá vỡ; có restore drill định kỳ.",
        "PII tối thiểu; không log password, token, QR secret hay toàn bộ thông tin nhạy cảm.",
        "Index theo truy vấn thật: published content, user progress, due review, order status và audit entity.",
    ])

    page_break(doc)
    doc.add_heading("05 · Kế hoạch nội dung và audio", level=1)
    add_table(doc, ["Lộ trình", "10–12 bài đầu", "Ví dụ tình huống"], [
        ["Văn phòng & hành chính", "Chào hỏi; nhận việc; lịch họp; email; tiến độ; xác nhận; báo cáo; xin hỗ trợ", "请问这个项目的进度如何？"],
        ["Nhà máy & sản xuất", "Ca làm; máy móc; nguyên liệu; thao tác; chất lượng; sự cố; bàn giao; an toàn cơ bản", "这台设备需要检查。"],
        ["Kho vận & logistics", "Nhập kho; xuất kho; tồn; mã hàng; đóng gói; giao nhận; trễ chuyến; đối chiếu", "请核对一下库存数量。"],
        ["Bán hàng & CSKH · sau", "Nhu cầu; báo giá; lịch hẹn; phản hồi; đổi trả; chăm sóc", "我来确认一下您的需求。"],
    ], widths=[2300, 4800, 2260])
    doc.add_heading("Quy trình audio tiết kiệm", level=2)
    add_numbered(doc, [
        "Khóa script sau reviewer; không thu khi nội dung còn thay đổi.",
        "Thu batch 30–50 câu trong phòng yên, cùng mic và khoảng cách; lưu master WAV và bản phát web.",
        "Chuẩn hóa tên file theo vocabulary_id/lesson_id; upload object storage; kiểm tra duration và checksum.",
        "QA ngẫu nhiên 20% và toàn bộ câu có số, thuật ngữ hoặc tên riêng.",
        "Nếu dùng TTS mã nguồn mở/thuê dịch vụ, phải gắn nguồn, giấy phép và chất lượng giọng; không cần speech scoring ở V1.",
    ])
    add_callout(doc, "Nguồn chuyên môn", "Bạn không cần tự biết mọi ngành. Hãy phỏng vấn 3–5 người đang làm mỗi ngành, xin tình huống lặp lại và câu họ thật sự dùng; sau đó để reviewer tiếng Trung chỉnh ngôn ngữ. Đây là vòng lặp nội dung đáng đầu tư nhất.", fill=MINT, icon="NỘI DUNG")

    page_break(doc)
    doc.add_heading("Đo lường, chất lượng và quyết định Flutter", level=1)
    doc.add_heading("Event tối thiểu", level=2)
    add_table(doc, ["Event", "Khi ghi", "Mục đích"], [
        ["signup_completed", "Tạo tài khoản xong", "Đo đầu phễu"],
        ["course_selected", "Chọn lộ trình đầu", "Activation step 1"],
        ["lesson_started / completed", "Mở / hoàn thành bài", "Completion và drop-off"],
        ["practice_answered", "Đánh giá một từ", "Độ khó và chất lượng SRS"],
        ["vip_checkout_started", "Tạo payment_order", "Ý định mua"],
        ["payment_confirmed", "Backend xác nhận", "Conversion thật"],
        ["subscription_activated", "Transaction thành công", "Kiểm soát payment → entitlement"],
    ], widths=[2700, 3200, 3460])
    doc.add_heading("North-star và guardrails", level=2)
    add_bullets(doc, [
        "North-star: số người hoàn thành ≥3 buổi học hữu ích mỗi tuần.",
        "Activation: chọn lộ trình + hoàn thành bài đầu trong 24 giờ.",
        "Retention: W1/W4 của người đã hoàn thành bài đầu, không chỉ số đăng ký.",
        "Conversion: active VIP / người học đã xem ít nhất 2 bài; tách theo nguồn vào.",
        "Guardrails: lỗi payment, webhook trùng, audio lỗi, nội dung bị report, thời gian phản hồi support.",
    ])
    doc.add_heading("Flutter gate", level=2)
    add_table(doc, ["Điều kiện", "Ngưỡng đề xuất", "Lý do"], [
        ["Quy mô", "≥300 WAU liên tiếp 4 tuần", "Đủ tín hiệu để đầu tư client thứ hai"],
        ["Retention", "W4 ≥25% trong nhóm activated", "Người học quay lại vì giá trị, không vì quảng cáo"],
        ["Doanh thu", "≥50 VIP active hoặc conversion ≥8%", "Có khả năng tài trợ vận hành app"],
        ["Nhu cầu native", "Có yêu cầu push/offline từ ≥20% người dùng phỏng vấn", "Flutter giải quyết vấn đề thật"],
        ["Backend", "API versioned + auth + sync test ổn", "Không xây app trên nền nghiệp vụ còn đổi liên tục"],
    ], widths=[2100, 2900, 4360])

    page_break(doc)
    doc.add_heading("Tính năng tương lai — xếp theo tín hiệu", level=1)
    add_table(doc, ["Thời điểm", "Tính năng", "Chỉ làm khi"], [
        ["Sau V1", "Bookmark, ghi chú cá nhân, lịch học, nhắc ôn", "Có người quay lại nhưng quên lịch"],
        ["Sau V1", "Download audio / PWA offline nhẹ", "Mạng yếu là lý do bỏ học"],
        ["Sau retention", "Flutter + push notification", "Đạt Flutter gate"],
        ["Sau content scale", "AI gợi ý ví dụ, kiểm tra trùng, tagging", "Admin CMS đã có reviewer và audit"],
        ["Sau PMF", "Luyện nói và feedback phát âm", "Nhiều người yêu cầu; có ngân sách model/audio"],
        ["B2B", "Lộ trình doanh nghiệp, team dashboard, seat billing", "Có ít nhất 3 khách hàng thử nghiệm"],
        ["Creator", "Cộng tác viên soạn bài, revenue share", "Quy trình review đủ chặt để mở rộng"],
    ], widths=[2100, 3700, 3560])
    add_callout(doc, "Không đưa vào backlog chỉ vì nghe hay", "Mỗi tính năng tương lai phải gắn với một metric hoặc một vấn đề quan sát được. Nếu không có tín hiệu, ưu tiên thêm bài học tốt và cải thiện vòng lặp quay lại.", fill=GOLD_LIGHT, icon="NGUYÊN TẮC")

    doc.add_heading("Kiểm tra trước public beta", level=2)
    add_table(doc, ["Nhóm", "Checklist nghiệm thu"], [
        ["Sản phẩm", "Onboarding, resume, khóa VIP, empty/error/loading state, mobile 390 px"],
        ["Nội dung", "30 bài duyệt; audio không 404; thuật ngữ có nguồn/người kiểm tra"],
        ["Thanh toán", "Webhook retry; trùng giao dịch; sai tiền; giao dịch muộn; đối soát"],
        ["Bảo mật", "RBAC; rate limit; secret rotation; dependency scan; backup/restore"],
        ["Pháp lý/vận hành", "Thông tin chủ thể; điều khoản; bảo mật; hoàn tiền; hỗ trợ; log consent"],
        ["Đo lường", "Event có user/order IDs đúng; dashboard activation/retention/conversion"],
    ], widths=[2100, 7260])

    page_break(doc)
    doc.add_heading("Kế hoạch 14 ngày tới", level=1)
    add_table(doc, ["Ngày", "Bạn A · Product/Frontend", "Bạn B · Backend/Ops", "Đầu ra chung"], [
        ["1–2", "Chốt tên, navigation, state màn hình", "Thiết lập env, PostgreSQL, migration", "Repo chạy ổn + database local"],
        ["3–4", "Form đăng nhập/hồ sơ", "Auth session + users", "Tài khoản thật vào dashboard"],
        ["5–6", "Nối course/lesson UI", "Query published content + seed", "Không còn dữ liệu khóa học hard-code"],
        ["7", "Test mobile, loading/error", "Logging + backup script", "Mốc demo nội bộ 1"],
        ["8–10", "Màn Admin CRUD cơ bản", "RBAC + course/module/lesson API", "Tạo bài không sửa DB tay"],
        ["11–12", "Soạn 3 bài mẫu theo template", "Content version + audit", "Luồng draft → preview"],
        ["13–14", "Usability test 3–5 người", "Fix dữ liệu, deploy staging", "Danh sách vấn đề có mức độ ưu tiên"],
    ], widths=[1000, 2800, 2800, 2760])
    add_callout(doc, "Điểm dừng sau 14 ngày", "Nếu người thử không hiểu phải học gì tiếp theo, dừng mở rộng backend và sửa onboarding/dashboard trước. Nếu họ hiểu nhưng nội dung chưa hữu ích, ưu tiên phỏng vấn và reviewer. Nếu cả hai ổn, mới tăng tốc Admin và VIP.", fill=MINT, icon="GATE")

    doc.add_heading("Tiêu chí chấp nhận của prototype hiện tại", level=2)
    add_bullets(doc, [
        "Build production thành công; TypeScript và ESLint không có lỗi.",
        "Tuyến /, /courses, /learn/[slug], /practice, /vip và /admin render đúng.",
        "Bản laptop và điện thoại giữ đúng thứ tự ưu tiên của phương án 3.",
        "Filter có trạng thái truy cập; ôn tập không cho trả lời trước khi mở thẻ; audio chưa có được disabled rõ ràng.",
        "Tệp design-qa.md ghi nhận P0/P1/P2 và kết quả cuối passed.",
    ])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("HẾT · HANZIWORK PRODUCT BLUEPRINT v0.2")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(FOREST)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_document())
